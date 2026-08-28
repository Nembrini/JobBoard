"""Estrazione del testo grezzo da un CV in PDF o DOCX.

Due estrattori PDF invece di uno perche' sbagliano in modi diversi:

* ``pypdfium2`` e' velocissimo (~9 ms) ma restituisce il testo nell'ordine del
  content stream: su un CV a due colonne le intreccia riga per riga. E su font
  con subset incompleto **perde glifi in silenzio**, tipicamente le cifre.
* ``pdfplumber`` ricostruisce il layout dalle coordinate, gestisce le colonne e
  recupera glifi che il primo scarta, al prezzo di ~160 ms.

Girano **entrambi, sempre**, e si tiene quello che ha conservato piu' contenuto
(vedi :func:`_pick_best`). Il primo tentativo di questo modulo usava il secondo
solo come fallback quando il primo "sembrava andato male": su un CV reale
``pypdfium2`` ha prodotto un testo di qualita' 0.98 — apparentemente perfetto —
avendo perso una cifra. Un'euristica che misura la somiglianza a prosa non puo'
accorgersi di un contenuto mancante.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

log = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx"})

#: Sotto questa soglia il PDF e' quasi certamente una scansione: non c'e' testo
#: da estrarre, servirebbe un OCR che questo progetto non fa.
_MIN_USABLE_CHARS = 200


class ExtractionError(RuntimeError):
    pass


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    #: Quale estrattore ha prodotto il risultato tenuto.
    method: Literal["pypdfium2", "pdfplumber", "python-docx"]
    pages: int | None
    language: str | None
    source_name: str
    #: Cifre che l'estrattore scartato aveva trovato e questo no.
    #: Non un errore: e' il prezzo consapevole di preferire un testo leggibile.
    #: Vanno mostrate in revisione, perche' sono i Result dei bullet ACR e
    #: sparirebbero altrimenti senza che nessuno se ne accorga.
    possibly_lost_numbers: tuple[str, ...] = ()

    @property
    def char_count(self) -> int:
        return len(self.text)


def extract(path: Path) -> ExtractedDocument:
    """Estrae il testo da un CV, scegliendo l'estrattore in base al formato."""
    if not path.exists():
        raise ExtractionError(f"file non trovato: {path}")
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ExtractionError(
            f"formato non supportato: {suffix or '(nessuna estensione)'}. "
            f"Attesi: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )

    doc = _extract_docx(path) if suffix == ".docx" else _extract_pdf(path)

    if doc.char_count < _MIN_USABLE_CHARS:
        raise ExtractionError(
            f"estratti solo {doc.char_count} caratteri da {path.name}. "
            "Probabilmente e' una scansione o un PDF di sole immagini: "
            "serve un originale con testo selezionabile."
        )
    return doc


# --- PDF ---------------------------------------------------------------------


@dataclass(frozen=True)
class _Candidate:
    method: Literal["pypdfium2", "pdfplumber"]
    text: str
    pages: int


def _extract_pdf(path: Path) -> ExtractedDocument:
    """Esegue entrambi gli estrattori e tiene quello che conserva piu' contenuto.

    Girano sempre tutti e due, anche quando il primo sembra riuscito: su un CV
    reale ``pypdfium2`` ha prodotto un testo dall'aspetto perfetto (qualita'
    0.98) **perdendo silenziosamente una cifra** — "circa 239 robot" diventava
    "circa robot". Un'euristica che misura solo quanto il testo somiglia a
    prosa non puo' accorgersene.

    Il costo e' irrilevante: un CV e' 1-2 pagine, pdfplumber impiega ~160 ms, e
    l'estrazione gira una volta per caricamento, non per annuncio.
    """
    fast_text, pages = _pdf_pypdfium2(path)
    candidates = [_Candidate("pypdfium2", fast_text, pages)]

    try:
        slow_text, slow_pages = _pdf_pdfplumber(path)
        candidates.append(_Candidate("pdfplumber", slow_text, slow_pages))
    except Exception as exc:  # pragma: no cover - dipende dal file
        log.warning("pdfplumber ha fallito su %s: %s", path.name, exc)

    best = _pick_best(candidates)

    # Cifre che il candidato scartato aveva e il vincitore no.
    kept = _numeric_tokens(best.text)
    lost = sorted({n for c in candidates if c is not best for n in _numeric_tokens(c.text)} - kept)
    if lost:
        log.warning(
            "%s: %s non contiene le cifre %s, trovate dall'altro estrattore",
            path.name,
            best.method,
            ", ".join(lost),
        )
    return _finish(best.text, best.method, best.pages, path, tuple(lost))


#: Sotto questa frazione di parole di lunghezza plausibile, l'estrazione ha
#: fuso le parole fra loro. Misurato su un CV reale: pypdfium2 1.00,
#: pdfplumber 0.73 sullo stesso file.
_MIN_WORD_INTEGRITY = 0.90

#: Oltre questa lunghezza una sequenza di lettere non e' una parola ma parole
#: incollate. L'italiano supera i 18 caratteri molto di rado.
_MAX_PLAUSIBLE_WORD = 18


def _pick_best(candidates: list[_Candidate]) -> _Candidate:
    """Sceglie l'estrazione che ha perso meno contenuto, in ordine di gravita'.

    I due estrattori sbagliano in modi di gravita' molto diversa, quindi la
    scelta e' gerarchica e non una somma di punteggi:

    1. **Confini di parola.** ``pdfplumber`` in modalita' layout, su certi font,
       restituisce ``macchinedatagliolaser`` invece di ``macchine da taglio
       laser``. E' il guasto peggiore: rende il testo inutilizzabile sia per
       l'LLM sia per il matching. Chi lo subisce viene escluso.
    2. **Cifre.** Fra i candidati leggibili vince chi ne conserva di piu':
       ``pypdfium2`` su font con subset incompleto le perde in silenzio, e le
       cifre sono i Result dei bullet ACR.
    3. **Qualita' della prosa**, solo per rompere la parita'.

    Una prima versione usava solo il criterio 2 e sceglieva un testo con tutte
    le parole fuse pur di recuperare un numero: barattare un documento
    leggibile per una cifra e' esattamente il contrario di quel che serve.
    """
    if len(candidates) == 1:
        return candidates[0]

    readable = [c for c in candidates if _word_integrity(c.text) >= _MIN_WORD_INTEGRITY]
    if not readable:
        # Nessuno e' pulito: si tiene il meno compromesso, ma vale la pena saperlo.
        worst_case = max(candidates, key=lambda c: _word_integrity(c.text))
        log.warning(
            "nessuna estrazione ha confini di parola affidabili "
            "(migliore: %s a %.2f); il profilo estratto va controllato a mano",
            worst_case.method,
            _word_integrity(worst_case.text),
        )
        return worst_case

    return max(readable, key=lambda c: (len(_numeric_tokens(c.text)), _quality(c.text)))


def _word_integrity(text: str) -> float:
    """Frazione di parole di lunghezza plausibile.

    Vale 1.0 su un testo normale e crolla quando l'estrattore perde gli spazi:
    non esiste un altro modo semplice di accorgersene, perche' un testo con le
    parole fuse ha righe lunghe e tante lettere, quindi supera indenne ogni
    euristica basata sulla forma.
    """
    words = re.findall(r"[^\W\d_]+", text, flags=re.UNICODE)
    if not words:
        return 0.0
    return sum(1 for w in words if len(w) <= _MAX_PLAUSIBLE_WORD) / len(words)


def _numeric_tokens(text: str) -> set[str]:
    """Token numerici distinti: percentuali, anni, quantita'.

    Il token deve terminare con una cifra (o con ``%``): senza questo vincolo la
    punteggiatura di fine frase entrerebbe nel token e ``2024,`` risulterebbe
    diverso da ``2024``, falsando il confronto fra estrattori.
    """
    return set(re.findall(r"\d(?:[\d.,]*\d)?%?", text))


def _pdf_pypdfium2(path: Path) -> tuple[str, int]:
    import pypdfium2

    pdf = pypdfium2.PdfDocument(str(path))
    try:
        pages = len(pdf)
        chunks = []
        for i in range(pages):
            page = pdf[i]
            textpage = page.get_textpage()
            chunks.append(textpage.get_text_bounded())
            textpage.close()
            page.close()
        return _clean("\n".join(chunks)), pages
    finally:
        pdf.close()


def _pdf_pdfplumber(path: Path) -> tuple[str, int]:
    import pdfplumber

    with pdfplumber.open(str(path)) as pdf:
        chunks = [p.extract_text(layout=True) or "" for p in pdf.pages]
        return _clean("\n".join(chunks)), len(pdf.pages)


# --- DOCX --------------------------------------------------------------------


def _extract_docx(path: Path) -> ExtractedDocument:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]

    # Molti CV in Word mettono le date o le skill dentro tabelle invisibili:
    # ignorarle perderebbe interi blocchi di contenuto.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(dict.fromkeys(cells)))

    return _finish(_clean("\n".join(parts)), "python-docx", None, path)


# --- utilita' ----------------------------------------------------------------


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    # I PDF spezzano le parole a fine riga: "svilup-\npato" -> "sviluppato".
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()


def _quality(text: str) -> float:
    """Punteggio euristico di quanto un'estrazione sembra riuscita.

    Non misura la correttezza — impossibile senza il file originale — ma
    distingue un testo plausibile da colonne intrecciate o caratteri persi.
    """
    if not text:
        return 0.0
    lines = [ln for ln in text.split("\n") if ln.strip()]
    if not lines:
        return 0.0
    avg_len = sum(len(ln) for ln in lines) / len(lines)
    letters = sum(c.isalpha() or c.isspace() for c in text) / len(text)
    # Righe di 1-3 caratteri in massa sono la firma di un layout a colonne
    # letto per coordinate sbagliate.
    fragments = sum(1 for ln in lines if len(ln) <= 3) / len(lines)
    return (min(avg_len, 80) / 80) * 0.5 + letters * 0.4 + (1 - fragments) * 0.1


def _detect_language(text: str) -> str | None:
    try:
        import py3langid

        code, _ = py3langid.classify(text[:4000])
        return str(code)
    except Exception:  # pragma: no cover - la lingua non e' mai bloccante
        return None


def _finish(
    text: str,
    method: Literal["pypdfium2", "pdfplumber", "python-docx"],
    pages: int | None,
    path: Path,
    possibly_lost_numbers: tuple[str, ...] = (),
) -> ExtractedDocument:
    return ExtractedDocument(
        text=text,
        method=method,
        pages=pages,
        language=_detect_language(text) if text else None,
        source_name=path.name,
        possibly_lost_numbers=possibly_lost_numbers,
    )
