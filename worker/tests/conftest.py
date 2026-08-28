"""Fixture condivise.

I PDF di prova vengono generati con Playwright invece di essere committati come
binari: restano leggibili nel diff, e il generatore e' lo stesso motore che in
Fase 6 impaginera' i CV veri, quindi il test esercita davvero quel percorso.
"""

from __future__ import annotations

from pathlib import Path

import pytest

_SINGLE_COLUMN = """
<h1>Filippo Nembrini</h1>
<p>Software Developer &middot; Milano, Italia &middot; filippo@example.com</p>
<h2>Experience</h2>
<h3>Backend Developer, Acme Srl (2022-01 - presente)</h3>
<ul>
  <li>Progettato e realizzato un servizio di fatturazione in Python e FastAPI,
      riducendo il tempo di elaborazione mensile da sei ore a venti minuti.</li>
  <li>Migrato il database da MySQL a PostgreSQL senza interruzioni di servizio,
      su un dataset di quaranta milioni di righe.</li>
</ul>
<h3>Junior Developer, Globex SpA (2020-09 - 2021-12)</h3>
<ul>
  <li>Sviluppate integrazioni REST con tre fornitori esterni, coperte da test
      automatici che hanno portato la copertura dal quaranta all'ottanta percento.</li>
</ul>
<h2>Skills</h2>
<p>Python, FastAPI, PostgreSQL, Docker, React, TypeScript, Git</p>
<h2>Education</h2>
<p>Laurea Triennale in Informatica, Universita' degli Studi di Milano, 2020</p>
"""

_TWO_COLUMN = """
<div style="column-count: 2; column-gap: 30px;">
  <h1>Filippo Nembrini</h1>
  <p>Software Developer &middot; Milano</p>
  <h2>Experience</h2>
  <p>Backend Developer presso Acme Srl, dal gennaio duemilaventidue a oggi.
     Progettato un servizio di fatturazione in Python e FastAPI che ha ridotto
     il tempo di elaborazione mensile da sei ore a venti minuti complessivi.</p>
  <p>Junior Developer presso Globex SpA, da settembre duemilaventi a dicembre
     duemilaventuno. Sviluppate integrazioni REST con tre fornitori esterni,
     coperte da test automatici scritti con pytest e mantenute nel tempo.</p>
  <h2>Skills</h2>
  <p>Python, FastAPI, PostgreSQL, Docker, React, TypeScript, Git, Linux</p>
</div>
"""

_PAGE = """<!doctype html><html><head><meta charset="utf-8"><style>
body {{ font-family: Georgia, serif; font-size: 11pt; line-height: 1.4; margin: 25mm; }}
h1 {{ font-size: 20pt; margin: 0 0 4pt; }}
h2 {{ font-size: 12pt; margin: 14pt 0 4pt; border-bottom: 1px solid #999; }}
h3 {{ font-size: 11pt; margin: 8pt 0 2pt; }}
</style></head><body>{body}</body></html>"""


def _render(html: str, target: Path) -> Path:
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="load")
        page.pdf(path=str(target), format="A4", print_background=True)
        browser.close()
    return target


@pytest.fixture(scope="session")
def single_column_cv(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """CV a colonna singola: il caso che l'estrattore veloce deve gestire da solo."""
    target = tmp_path_factory.mktemp("cv") / "single.pdf"
    return _render(_PAGE.format(body=_SINGLE_COLUMN), target)


@pytest.fixture(scope="session")
def two_column_cv(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """CV a due colonne: il caso che fa inciampare l'estrazione per content stream."""
    target = tmp_path_factory.mktemp("cv") / "two-column.pdf"
    return _render(_PAGE.format(body=_TWO_COLUMN), target)


@pytest.fixture(scope="session")
def docx_cv(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """CV in Word, con una tabella: i CV veri ci mettono dentro date e skill."""
    import docx

    target = tmp_path_factory.mktemp("cv") / "cv.docx"
    d = docx.Document()
    d.add_heading("Filippo Nembrini", level=1)
    d.add_paragraph("Software Developer - Milano, Italia - filippo@example.com")
    d.add_heading("Experience", level=2)
    d.add_paragraph(
        "Backend Developer presso Acme Srl. Progettato un servizio di fatturazione "
        "in Python e FastAPI, riducendo il tempo di elaborazione da sei ore a venti minuti."
    )
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Periodo"
    table.cell(0, 1).text = "2022-01 - presente"
    table.cell(1, 0).text = "Stack"
    table.cell(1, 1).text = "Python, FastAPI, PostgreSQL, Docker"
    d.save(str(target))
    return target
